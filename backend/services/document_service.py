import json
import shutil
import subprocess

from config import Config
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def create_letter_doc(folder_path, content):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    lines = [line.rstrip() for line in content.split("\n")]

    for i, line in enumerate(lines, start=1):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

        if i in (6, 7):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if i in [5, 7, 9, len(lines) - 2]:
            doc.add_paragraph("")

    doc.save(folder_path / "Lettre de motivation César Fuentes.docx")


def create_overrides_json(folder_path, results):
    overrides_data = {}

    for key, value in results.items():
        if key == "letter":
            continue

        try:
            parsed = json.loads(value)
            if isinstance(parsed, dict) and key in parsed:
                overrides_data[key] = parsed[key]
            else:
                overrides_data[key] = parsed
        except json.JSONDecodeError:
            print(f"Warning: {key} is not valid JSON, storing as string")
            overrides_data[key] = value

    output_path = folder_path / "overrides.json"
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(overrides_data, f, ensure_ascii=False, indent=2)

    print(f"Overrides JSON created: {output_path}")


def generate_resume_pdf(folder_path, company, date):
    try:
        Config.RESUME_OVERRIDES_DIR.mkdir(parents=True, exist_ok=True)

        overrides_json_path = folder_path / "overrides.json"

        if overrides_json_path.exists():
            with open(overrides_json_path, "r", encoding="utf-8") as f:
                overrides_data = json.load(f)

            for key, value in overrides_data.items():
                override_path = Config.RESUME_OVERRIDES_DIR / f"{key}.json"
                with open(override_path, "w", encoding="utf-8") as f:
                    json.dump(value, f, ensure_ascii=False, indent=2)
                print(f"Created override: {key}.json")

        try:
            output_pdf = folder_path / f"CV César Fuentes.pdf"

            result = subprocess.run(
                [
                    "node",
                    str(Config.RESUME_SCRIPT),
                    str(output_pdf),
                ],
                capture_output=True,
                text=True,
                check=True,
            )

            print(result.stdout)
            print(f"Resume PDF generated: {output_pdf}")
            return str(output_pdf)

        finally:
            if Config.RESUME_OVERRIDES_DIR.exists():
                shutil.rmtree(Config.RESUME_OVERRIDES_DIR)
                print(f"Cleaned up resume_overrides folder")

    except subprocess.CalledProcessError as e:
        print(f"Error generating PDF: {e.stderr}")
        raise
    except Exception as e:
        print(f"Error: {e}")
        raise
