import json
import os
import subprocess
from datetime import datetime
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from dotenv import load_dotenv
from flask import Flask, jsonify, request
from flask_cors import CORS
from openai import OpenAI
from openpyxl.styles import Font

app = Flask(__name__)
CORS(
    app,
    resources={
        r"/api/*": {
            "origins": "*",
            "methods": ["GET", "POST", "OPTIONS"],
            "allow_headers": ["Content-Type"],
        }
    },
)

load_dotenv()
API_KEY = os.getenv("API_KEY")
BASE_DIR = Path(__file__).parent.parent
CANDIDATURES_DIR = BASE_DIR / "Applications"
PROMPTS_DIR = BASE_DIR / "prompts"
EXCEL_FILE = BASE_DIR / "Recherche Janvier 2026.xlsx"
RESUME_DIR = BASE_DIR / "resume"
RESUME_JSON = RESUME_DIR / "public" / "resume.json"
RESUME_SCRIPT = RESUME_DIR / "exportToPDF.js"
PHOTO_PATH = RESUME_DIR / "src" / "assets" / "photo.jpeg"
DEBUG_MODE = False

client = OpenAI(api_key=API_KEY)


def query_openai(prompt):
    response = client.chat.completions.create(
        model="gpt-4.1", messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()


def get_prompts():
    prompts = {}
    prompt_map = {
        "prompt_1.txt": "HARD_SKILLS",
        "prompt_2.txt": "SOFT_SKILLS",
        "prompt_3.txt": "LETTER",
    }

    for filename, key in prompt_map.items():
        filepath = PROMPTS_DIR / filename
        if filepath.exists():
            with open(filepath, "r", encoding="utf-8") as f:
                prompts[key] = f.read().strip()

    return prompts


def extract_job_info(content):
    prompt = f"""
À partir du texte de page web ci-dessous, retourne uniquement en JSON, sans aucune autre information:
{{
  "company": "Nom de la société qui poste l'offre d'emploi",
  "job_title": "Intitulé du poste",
  "platform": "Plateforme d'offres d'emploi, en un mot",
  "job_description": "Contenu de l'offre d'emploi et du profil recherché"
}}
Offre d'emploi:
{content}
"""
    response = query_openai(prompt)

    try:
        return json.loads(response)
    except json.JSONDecodeError:
        import re

        match = re.search(r"\{.*\}", response, re.DOTALL)
        if match:
            json_str = match.group(0)
            try:
                return json.loads(json_str)
            except Exception as e:
                print(f"❌ Still not valid JSON: {e}")
        raise ValueError(f"Invalid JSON returned by model: {response}")


def add_to_excel(company, platform, job_title, url, date):
    """Add entry to Excel, format new row, then sort all rows by status (column A)"""
    try:
        wb = openpyxl.load_workbook(EXCEL_FILE)
        ws = wb["suivi"]

        new_row_data = ["A faire", "CDI", company, platform, job_title, url, date, ""]
        ws.append(new_row_data)
        new_row_idx = ws.max_row

        for col in range(1, len(new_row_data) + 1):
            ws.cell(row=new_row_idx, column=col).font = Font(name="Roboto", size=10)

        for cell in ws[1]:
            cell.font = Font(name="Roboto", size=11)

        data = list(ws.iter_rows(min_row=2, values_only=True))
        data_sorted = sorted(data, key=lambda x: x[0] or "")

        for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
            for cell in row:
                cell.value = None

        for r_idx, row_data in enumerate(data_sorted, start=2):
            for c_idx, value in enumerate(row_data, start=1):
                cell = ws.cell(row=r_idx, column=c_idx, value=value)
                cell.font = Font(name="Roboto", size=11)

        wb.save(EXCEL_FILE)
        wb.close()

    except Exception as e:
        print(f"Error adding to Excel: {e}")
        raise


def create_letter_doc(folder_path, content):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    lines = [line.rstrip() for line in content.split("\n")]

    for i, line in enumerate(lines, start=1):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

        if i in (6, 7):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if i in [5, 7, 9, len(lines) - 2]:
            doc.add_paragraph("")

    doc.save(folder_path / "Lettre de motivation César Fuentes.docx")


def create_skills_json(folder_path, results):
    """Create a JSON file with the skills data"""
    skills_data = {
        "hard_skills": {
            "main": [results["HARD_SKILLS"].split(",")[0].strip()],
            "secondary": (
                [
                    skill.strip()
                    for skill in results["HARD_SKILLS"].split(",")[1].split(",")
                    if skill.strip()
                ]
                if "," in results["HARD_SKILLS"]
                and len(results["HARD_SKILLS"].split(",")) > 1
                else []
            ),
            "environmnent_and_tools": (
                [
                    skill.strip()
                    for skill in results["HARD_SKILLS"].split(",")[2].split(",")
                    if skill.strip()
                ]
                if "," in results["HARD_SKILLS"]
                and len(results["HARD_SKILLS"].split(",")) > 2
                else []
            ),
        },
        "soft_skills": [
            skill.strip()
            for skill in results["SOFT_SKILLS"].split("\n")
            if skill.strip()
        ],
    }

    with open(folder_path / "skills.json", "w", encoding="utf-8") as f:
        json.dump(skills_data, f, ensure_ascii=False, indent=2)

    print(f"✅ Skills JSON created: {folder_path / 'skills.json'}")


def generate_resume_pdf(folder_path, company, date):
    """Generate a PDF resume with custom skills from skills.json"""
    try:
        # Create resume_overrides directory in public folder
        overrides_dir = RESUME_DIR / "public" / "resume_overrides"
        overrides_dir.mkdir(exist_ok=True)

        # Path to skills.json
        skills_json_path = folder_path / "skills.json"

        # Copy skills.json to overrides folder if it exists
        if skills_json_path.exists():
            import shutil

            with open(skills_json_path, "r", encoding="utf-8") as f:
                skills_data = json.load(f)

            # Copy hard_skills if exists
            if "hard_skills" in skills_data:
                hard_skills_path = overrides_dir / "hard_skills.json"
                with open(hard_skills_path, "w", encoding="utf-8") as f:
                    json.dump(
                        skills_data["hard_skills"], f, ensure_ascii=False, indent=2
                    )
                print(f"✅ Copied hard_skills.json to overrides folder")

            # Copy soft_skills if exists
            if "soft_skills" in skills_data:
                soft_skills_path = overrides_dir / "soft_skills.json"
                with open(soft_skills_path, "w", encoding="utf-8") as f:
                    json.dump(
                        skills_data["soft_skills"], f, ensure_ascii=False, indent=2
                    )
                print(f"✅ Copied soft_skills.json to overrides folder")

        try:
            # Output PDF path
            output_pdf = folder_path / f"CV César Fuentes.pdf"

            # Call Node.js script to generate PDF
            result = subprocess.run(
                [
                    "node",
                    str(RESUME_SCRIPT),
                    str(output_pdf),
                ],
                capture_output=True,
                text=True,
                check=True,
            )

            print(result.stdout)
            print(f"✅ Resume PDF generated: {output_pdf}")
            return str(output_pdf)

        finally:
            # Clean up: remove all files from overrides folder
            if overrides_dir.exists():
                import shutil

                shutil.rmtree(overrides_dir)
                print(f"✅ Cleaned up resume_overrides folder")

    except subprocess.CalledProcessError as e:
        print(f"❌ Error generating PDF: {e.stderr}")
        raise
    except Exception as e:
        print(f"❌ Error: {e}")
        raise


def generate_job_documents(company, job_title, job_content):
    prompts = get_prompts()

    results = {}
    for key, prompt in prompts.items():
        results[key] = query_openai(f"{prompt}\n\n{job_content}")

    today = datetime.now()
    formatted_date = today.strftime("%Y-%m-%d")
    folder_name = f"{company} - {formatted_date}"
    folder_path = CANDIDATURES_DIR / folder_name
    folder_path.mkdir(parents=True, exist_ok=True)

    # Create letter
    create_letter_doc(folder_path, results["LETTER"])

    # Create skills JSON (instead of Word doc)
    create_skills_json(folder_path, results)

    # Generate PDF resume with custom skills
    generate_resume_pdf(folder_path, company, formatted_date)

    return str(folder_path.absolute())


@app.route("/api/job", methods=["POST", "OPTIONS"])
def handle_job():
    if request.method == "OPTIONS":
        return "", 204

    try:
        data = request.json
        extracted = extract_job_info(data["content"])
        company = extracted.get("company", "Inconnue").title()
        job_title = extracted.get("job_title", data.get("title", "Poste")).title()
        job_content = extracted.get("job_description", data.get("content", ""))
        platform = extracted.get("platform", "Hors Plateforme").title()

        today = datetime.now().strftime("%Y-%m-%d")
        add_to_excel(company, platform, job_title, data["url"], today)
        folder_path = generate_job_documents(company, job_title, job_content)

        return jsonify(
            {
                "status": "success",
                "folderPath": folder_path,
                "company": company,
                "jobTitle": job_title,
            }
        )
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


if __name__ == "__main__":
    CANDIDATURES_DIR.mkdir(exist_ok=True)
    PROMPTS_DIR.mkdir(exist_ok=True)

    print(f"Server starting...")
    print(f"Applications folder: {CANDIDATURES_DIR}")
    print(f"Prompts folder: {PROMPTS_DIR}")
    print(f"Excel file: {EXCEL_FILE}")
    print(f"Resume script: {RESUME_SCRIPT}")

    app.run(debug=True, port=5000)
