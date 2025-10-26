import traceback
from pathlib import Path

import openpyxl
from config import Config
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl.styles import Font


def add_to_sheet(company, platform, job_title, url, date):
    try:
        wb = openpyxl.load_workbook(Config.EXCEL_FILE)
        ws = wb["suivi"]

        new_row_data = ["A faire", "CDI", company, platform, job_title, url, date, ""]
        ws.append(new_row_data)
        new_row_idx = ws.max_row

        for col in range(1, len(new_row_data) + 1):
            ws.cell(row=new_row_idx, column=col).font = Font(name="Roboto", size=10)

        for cell in ws[1]:
            cell.font = Font(name="Roboto", size=11)

        data = list(ws.iter_rows(min_row=2, values_only=True))
        data_sorted = sorted(data, key=lambda x: x[0] or "")

        for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
            for cell in row:
                cell.value = None

        for r_idx, row_data in enumerate(data_sorted, start=2):
            for c_idx, value in enumerate(row_data, start=1):
                cell = ws.cell(row=r_idx, column=c_idx, value=value)
                cell.font = Font(name="Roboto", size=11)

        wb.save(Config.EXCEL_FILE)
        wb.close()

    except Exception as e:
        print(f"Error adding to Excel: {e}")
        raise


def make_letter(openai, folder_path):
    folder_path = Path(folder_path)

    try:
        letter_path = Config.PROMPTS_DIR / "letter.txt"
        with open(letter_path, "r", encoding="utf-8") as f:
            letter_prompt = f.read().strip()

        content = openai.query(letter_prompt, True, True)

        doc = Document()

        try:
            style = doc.styles["Normal"]
            font = style.font
            font.name = "Times New Roman"
            font.size = Pt(11)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.line_spacing = 1.0
        except Exception as e:
            print("[ERROR] Style config failed:", e)
            traceback.print_exc()

        lines = [line.rstrip() for line in content.split("\n")]

        for i, line in enumerate(lines, start=1):
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

            if i in (6, 7):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if i in [5, 7, 10, len(lines) - 2]:
                doc.add_paragraph("")

        output_path = folder_path / "Lettre de motivation César Fuentes.docx"
        doc.save(str(output_path))

    except Exception as e:
        traceback.print_exc()
